from docx import Document
from docx.shared import Inches
import pyttsx3

document = Document()

pyttsx3.speak("Please Enter Your Full Name")
full_name = input("Please Enter Your Full Name : ")

pyttsx3.speak("Please Enter Your First Name")
first_name = input("Please Enter Your First Name : ")
pyttsx3.speak(f"Hi {first_name} ")

pyttsx3.speak("Please Enter Your Last Name")
last_name = input("Please Enter Your Last Name : ")

pyttsx3.speak("Please Enter Your Date of Birth ")
date_of_birth = input("Please Enter Your Date of Birth : ")

pyttsx3.speak("Please Enter Your email")
email = input("Please Enter Your Email : ")

pyttsx3.speak("Please Enter Your phone number")
phone_number = input("Please Enter Your Phone Number : ")

pyttsx3.speak("Please Enter the name of your picture file ")
picture = input(
    "Please Enter the name of your picture or logo (Make sure the picture is in your folder): ")

pyttsx3.speak("Please Enter about you")
about = input("Please say me about You : ")

document.add_picture(
    picture,
    width=Inches(2.0)
)

document.add_heading("Basic Information")

document.add_paragraph(f"""

Full Name : {full_name}

Last Name : {first_name}

Full Name : {last_name}

Full Name : {date_of_birth}

Full Name : {email}

Full Name : {phone_number}

""")

document.add_heading("About Me")

document.add_paragraph(
    about
)

document.add_heading(" Work Experience ")

document.add_paragraph(input("Say me about your work experience: "))

document.save("cv.docx")
